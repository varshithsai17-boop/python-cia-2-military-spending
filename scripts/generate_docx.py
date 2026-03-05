"""
generate_docx.py — Generate Word Document with Marker & Color Customization Codes
==================================================================================
Creates 'docs/additional_codes.docx' containing all custom marker, color, and
styling code examples used in the analysis, formatted in Times New Roman 12pt.

Usage:
  python scripts/generate_docx.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(PROJECT, "docs")
os.makedirs(DOCS_DIR, exist_ok=True)

doc = Document()

# ── Default font ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ── Helper functions ─────────────────────────────────────────────────────────
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_code_block(code_text, description=""):
    if description:
        p = doc.add_paragraph()
        run = p.add_run(description)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_text(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('Additional Codes: Custom Markers & Colors', level=0)
for run in title.runs:
    run.font.name = 'Times New Roman'
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_text("Python for Business — CIA Project", bold=True)
add_text("Data: Global Military Expenditure (Wikipedia)")
add_text("Libraries: matplotlib, seaborn, numpy")
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — CHANGING BAR COLORS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("1. Changing Bar Colors", level=1)
add_text("You can pass a list of colors to the 'color' parameter in plt.bar().")

add_code_block(
"""import matplotlib.pyplot as plt

countries = ['United States', 'China', 'Russia', 'Germany', 'India']
spending = [997.0, 314.0, 149.0, 88.5, 86.1]

# Method 1 — Named colors
plt.bar(countries, spending, color=['red', 'blue', 'green', 'orange', 'purple'])
plt.title('Top 5 Military Spenders')
plt.ylabel('Spending (US$ Billion)')
plt.xticks(rotation=45, ha='right')
plt.tight_layout()
plt.show()""",
"Example 1.1 — Using named colors:")

add_code_block(
"""# Method 2 — Hex color codes
custom_colors = ['#e74c3c', '#3498db', '#2ecc71', '#f39c12', '#9b59b6']
plt.bar(countries, spending, color=custom_colors)
plt.title('Using Hex Color Codes')
plt.tight_layout()
plt.show()""",
"Example 1.2 — Using hex color codes:")

add_code_block(
"""# Method 3 — Colormap gradient
import numpy as np
colors = plt.cm.viridis(np.linspace(0.2, 0.9, len(countries)))
plt.bar(countries, spending, color=colors)
plt.title('Using Viridis Colormap')
plt.tight_layout()
plt.show()""",
"Example 1.3 — Using colormap gradient:")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — CHANGING LINE MARKERS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("2. Changing Line Markers", level=1)
add_text("Available matplotlib marker styles:")

markers_table = doc.add_table(rows=9, cols=2)
markers_table.style = 'Table Grid'
headers = markers_table.rows[0].cells
headers[0].text = 'Marker Code'
headers[1].text = 'Description'
marker_data = [
    ("'o'", "Circle"), ("'s'", "Square"), ("'^'", "Triangle Up"),
    ("'D'", "Diamond"), ("'v'", "Triangle Down"), ("'p'", "Pentagon"),
    ("'*'", "Star"), ("'h'", "Hexagon")
]
for i, (code, desc) in enumerate(marker_data, 1):
    markers_table.rows[i].cells[0].text = code
    markers_table.rows[i].cells[1].text = desc

doc.add_paragraph()

add_code_block(
"""# Using different markers on a line chart
x = range(5)
spending = [997.0, 314.0, 149.0, 88.5, 86.1]

plt.plot(x, spending, marker='o', color='red', markersize=10, label='Circle')
plt.plot(x, [s * 0.9 for s in spending], marker='s', color='blue',
         markersize=10, label='Square')
plt.plot(x, [s * 0.8 for s in spending], marker='^', color='green',
         markersize=10, label='Triangle')

plt.legend()
plt.title('Line Chart with Different Markers')
plt.show()""",
"Example 2.1 — Different markers on the same chart:")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — SCATTER PLOT MARKERS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("3. Scatter Plot — Custom Markers & Sizes", level=1)

add_code_block(
"""# Scatter with custom marker, size, and color
spending = [997.0, 314.0, 149.0, 88.5, 86.1]
gdp_pct  = [3.4, 1.7, 7.1, 1.9, 2.3]

plt.scatter(spending, gdp_pct,
           marker='D',              # Diamond marker
           s=200,                    # Size
           c=['#e74c3c', '#3498db', '#2ecc71', '#f39c12', '#9b59b6'],
           edgecolors='black',       # Border color
           linewidth=1.5,            # Border width
           alpha=0.8)                # Transparency

plt.title('Scatter Plot with Diamond Markers')
plt.xlabel('Spending (US$ Billion)')
plt.ylabel('% of GDP')
plt.show()""",
"Example 3.1 — Diamond markers with custom colors:")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — COLOR PALETTES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("4. Color Palettes (Matplotlib & Seaborn)", level=1)

add_text("Common matplotlib colormaps: viridis, plasma, inferno, magma, cividis, coolwarm, RdYlGn, Set1, Set2, Pastel1")

add_code_block(
"""import seaborn as sns
import numpy as np

fig, axes = plt.subplots(2, 2, figsize=(12, 8))
data = [997.0, 314.0, 149.0, 88.5, 86.1]
labels = ['US', 'CHN', 'RUS', 'GER', 'IND']

# Palette 1 — matplotlib viridis
axes[0,0].bar(labels, data, color=plt.cm.viridis(np.linspace(0.2, 0.9, 5)))
axes[0,0].set_title('Viridis')

# Palette 2 — matplotlib plasma
axes[0,1].bar(labels, data, color=plt.cm.plasma(np.linspace(0.2, 0.9, 5)))
axes[0,1].set_title('Plasma')

# Palette 3 — seaborn pastel
axes[1,0].bar(labels, data, color=sns.color_palette('pastel', 5))
axes[1,0].set_title('Seaborn Pastel')

# Palette 4 — seaborn Set2
axes[1,1].bar(labels, data, color=sns.color_palette('Set2', 5))
axes[1,1].set_title('Seaborn Set2')

plt.suptitle('Four Different Color Palettes', fontweight='bold')
plt.tight_layout()
plt.show()""",
"Example 4.1 — Comparing four different color palettes:")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — HATCH PATTERNS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("5. Hatch Patterns on Bar Charts", level=1)

add_code_block(
"""countries = ['US', 'CHN', 'RUS', 'GER', 'IND']
spending = [997.0, 314.0, 149.0, 88.5, 86.1]
colors = ['#e74c3c', '#3498db', '#2ecc71', '#f39c12', '#9b59b6']
hatches = ['/', '\\\\', 'x', '.', 'o']

bars = plt.bar(countries, spending, color=colors, edgecolor='black', linewidth=1.5)
for bar, hatch in zip(bars, hatches):
    bar.set_hatch(hatch)

plt.title('Bar Chart with Hatch Patterns')
plt.ylabel('Spending (US$ Billion)')
plt.show()

# Available hatch patterns: / \\\\ | - + x o O . *""",
"Example 5.1 — Adding hatch patterns to bars:")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PIE CHART CUSTOMIZATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("6. Pie Chart — Custom Colors & Explode", level=1)

add_code_block(
"""labels = ['US', 'China', 'Russia', 'Germany', 'India', 'Others']
sizes = [35.5, 11.2, 5.5, 3.2, 3.1, 41.5]
colors = ['#e74c3c', '#f39c12', '#2ecc71', '#3498db', '#9b59b6', '#95a5a6']
explode = [0.08, 0.03, 0.03, 0, 0, 0]  # explode US and China

plt.pie(sizes, labels=labels, colors=colors, explode=explode,
        autopct='%1.1f%%', shadow=True, startangle=140,
        textprops={'fontsize': 11, 'fontweight': 'bold'})
plt.title('Global Military Spending Share')
plt.show()""",
"Example 6.1 — Exploding slices with custom colors:")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — FONT & TITLE STYLING
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("7. Font & Title Styling", level=1)

add_code_block(
"""plt.bar(['US', 'China', 'Russia'], [997, 314, 149], color='steelblue')

# Title styling
plt.title('Military Spending',
          fontsize=18,
          fontweight='bold',
          fontstyle='italic',
          fontfamily='serif',
          color='darkblue')

# Axis label styling
plt.xlabel('Country', fontsize=14, fontfamily='sans-serif', color='gray')
plt.ylabel('US$ Billion', fontsize=14, fontfamily='sans-serif', color='gray')

# Tick styling
plt.xticks(fontsize=12, rotation=0)
plt.yticks(fontsize=10)

plt.show()""",
"Example 7.1 — Customising fonts, sizes, and styles:")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — LEGEND POSITIONING
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("8. Legend Positioning", level=1)

add_code_block(
"""x = range(5)
plt.plot(x, [997, 314, 149, 88, 86], 'o-r', label='SIPRI 2024')
plt.plot(x, [921, 251, 186, 107, 94], 's-b', label='IISS 2025')

# Legend options:
# loc: 'upper right', 'upper left', 'lower left', 'lower right',
#      'center', 'best' (auto)
# bbox_to_anchor: (x, y) for precise placement

plt.legend(loc='upper right',
           fontsize=12,
           frameon=True,
           shadow=True,
           fancybox=True,
           framealpha=0.9,
           edgecolor='black')

plt.title('Legend Positioning Example')
plt.show()""",
"Example 8.1 — Customising legend position and style:")

# ─── Save ────────────────────────────────────────────────────────────────────
output_path = os.path.join(DOCS_DIR, "additional_codes.docx")
doc.save(output_path)
print(f"Word document saved: {output_path}")
